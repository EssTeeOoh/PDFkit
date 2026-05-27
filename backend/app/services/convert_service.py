import os
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image, ImageFilter

from app.config import TESSERACT, POPPLER, OCR_DPI
from app.logger import get_logger

logger = get_logger(__name__)



# ── Text cleaning ────────────────────────────────────────
import re

# Map common CID codes that pdfplumber can't decode from embedded fonts
_CID_MAP = {
    # Bullets & dashes
    "127": "•",   "149": "•",   "183": "•",
    "150": "–",   "151": "—",   "173": "-",
    # Quotes
    "147": "\u201c", "148": "\u201d",   # " "
    "145": "\u2018", "146": "\u2019",   # ' '
    # Other common ones
    "160": " ",   "169": "©",   "174": "®",
    "176": "°",   "177": "±",   "215": "×",
    "8": "•",     "118": "✓",   "252": "✓",
    "167": "§",   "182": "¶",
}

_CID_RE = re.compile(r"\(cid:(\d+)\)")


def _clean_text(text: str) -> str:
    """
    Replace (cid:N) sequences with their likely Unicode equivalents,
    then strip any remaining unresolved ones.
    Also normalises other common encoding artifacts.
    """
    if not text:
        return text

    def _replace(m):
        return _CID_MAP.get(m.group(1), "")   # drop unknown cid codes

    text = _CID_RE.sub(_replace, text)

    # Normalize common ligatures that sometimes come through as garbage
    text = text.replace("\ufb01", "fi").replace("\ufb02", "fl")
    text = text.replace("\ufb03", "ffi").replace("\ufb04", "ffl")
    text = text.replace("\u0000", "")   # null bytes

    return text


# ── Document helpers ────────────────────────────────────
def _apply_margins(doc: Document):  # type: ignore[misc]
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)


# ── PDF type detection ──────────────────────────────────
def is_scanned_pdf(pdf_path: str) -> bool:
    """
    Returns True if the PDF contains no extractable text on any page.
    Logs per-page text length so you can see exactly what pdfplumber finds.
    """
    logger.debug(f"[is_scanned_pdf] Checking: {pdf_path}")
    try:
        with pdfplumber.open(pdf_path) as pdf:
            logger.debug(f"[is_scanned_pdf] Total pages: {len(pdf.pages)}")
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                char_count = len(text.strip()) if text else 0
                logger.debug(f"[is_scanned_pdf] Page {i+1}: {char_count} chars extracted")
                if char_count > 0:
                    logger.info(f"[is_scanned_pdf] Text found on page {i+1} → digital PDF")
                    return False
        logger.info("[is_scanned_pdf] No text found on any page → scanned PDF")
        return True
    except Exception as e:
        logger.error(f"[is_scanned_pdf] Error reading PDF: {e}")
        return True


# ── Image preprocessing ─────────────────────────────────
def preprocess_image(image: Image.Image, page_num: int) -> Image.Image:
    """
    Preprocess a page image before OCR:
    1. Convert to greyscale
    2. Upscale if too small
    3. Sharpen
    4. Binarize (threshold)
    Logs dimensions at each stage.
    """
    logger.debug(f"[preprocess] Page {page_num} original size: {image.size}, mode: {image.mode}")

    # 1. Greyscale
    image = image.convert("L")
    logger.debug(f"[preprocess] Page {page_num} after greyscale: {image.size}")

    # 2. Upscale if too small
    width, height = image.size
    if width < 2000:
        scale = 2000 / width
        new_size = (int(width * scale), int(height * scale))
        image = image.resize(new_size, Image.Resampling.LANCZOS)
        logger.debug(f"[preprocess] Page {page_num} upscaled to: {image.size}")
    else:
        logger.debug(f"[preprocess] Page {page_num} no upscale needed")

    # 3. Sharpen
    image = image.filter(ImageFilter.SHARPEN)

    # 4. Binarize
    image = image.point(lambda x: 0 if int(x) < 140 else 255, "1")  # type: ignore[arg-type]
    image = image.convert("L")
    logger.debug(f"[preprocess] Page {page_num} binarized, final size: {image.size}")

    return image


# ── Text-based PDF → DOCX ───────────────────────────────
def extract_text_pdf_to_docx(pdf_path: str, output_path: str, original_name: str):
    """Convert a text-based PDF to .docx using pdfplumber."""
    logger.info(f"[pdf-to-word] Starting TEXT extraction: {pdf_path}")
    doc = Document()
    _apply_margins(doc)

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        logger.info(f"[pdf-to-word] Pages to process: {total_pages}")

        for page_num, page in enumerate(pdf.pages):
            if page_num > 0:
                doc.add_page_break()

            text = _clean_text(page.extract_text(layout=True))
            char_count = len(text.strip()) if text else 0
            logger.debug(f"[pdf-to-word] Page {page_num+1}: {char_count} chars")

            if not text or not text.strip():
                logger.warning(f"[pdf-to-word] Page {page_num+1}: no text — inserting placeholder")
                doc.add_paragraph(f"[Page {page_num + 1} — no extractable text]")
                continue

            lines_written = 0
            for line in text.split("\n"):
                stripped = _clean_text(line.strip())
                if not stripped:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    continue
                p = doc.add_paragraph()
                is_heading = (
                    stripped.isupper() and 3 < len(stripped) < 80
                ) or (len(stripped) < 60 and stripped.endswith(":"))
                run = p.add_run(stripped)
                run.bold = is_heading
                run.font.size = Pt(13) if is_heading else Pt(11)
                p.paragraph_format.space_after = Pt(4)
                lines_written += 1

            logger.debug(f"[pdf-to-word] Page {page_num+1}: wrote {lines_written} lines")

            # Tables
            tables = page.extract_tables()
            logger.debug(f"[pdf-to-word] Page {page_num+1}: found {len(tables)} table(s)")
            for t_idx, table_data in enumerate(tables):
                if not table_data:
                    continue
                try:
                    rows = len(table_data)
                    cols = max(len(row) for row in table_data)
                    tbl = doc.add_table(rows=rows, cols=cols)
                    tbl.style = "Table Grid"
                    for r_idx, row in enumerate(table_data):
                        for c_idx, val in enumerate(row):
                            if c_idx < cols:
                                tbl.rows[r_idx].cells[c_idx].text = str(val or "")
                    doc.add_paragraph()
                    logger.debug(f"[pdf-to-word] Page {page_num+1}: table {t_idx+1} written ({rows}x{cols})")
                except Exception as e:
                    logger.warning(f"[pdf-to-word] Page {page_num+1}: table {t_idx+1} skipped — {e}")

    doc.save(output_path)

    # Verify output file was actually written and has content
    size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
    logger.info(f"[pdf-to-word] Saved: {output_path} ({size} bytes)")
    if size < 5000:
        logger.warning(f"[pdf-to-word] Output file is very small ({size} bytes) — may be empty")


# ── Scanned PDF → DOCX (OCR) ────────────────────────────
def ocr_pdf_to_docx(pdf_path: str, output_path: str, original_name: str):
    """
    Convert a scanned/image-based PDF to .docx using Tesseract OCR.
    Logs every stage so you can pinpoint exactly where it fails.
    """
    logger.info(f"[ocr] Starting OCR conversion: {pdf_path}")
    logger.info(f"[ocr] Tesseract path: {TESSERACT}")
    logger.info(f"[ocr] Poppler path:   {POPPLER}")
    logger.info(f"[ocr] OCR DPI:        {OCR_DPI}")

    pytesseract.pytesseract.tesseract_cmd = TESSERACT

    # Verify Tesseract is actually callable
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"[ocr] Tesseract version: {version}")
    except Exception as e:
        logger.error(f"[ocr] Tesseract not callable: {e}")
        raise RuntimeError(f"Tesseract is not working: {e}")

    TESS_CONFIG = "--oem 3 --psm 6"
    doc = Document()
    _apply_margins(doc)

    # Step 1: Convert PDF pages → images
    logger.info(f"[ocr] Converting PDF pages to images at {OCR_DPI} DPI...")
    try:
        images = convert_from_path(
            pdf_path,
            dpi=OCR_DPI,
            poppler_path=str(POPPLER) if POPPLER else None,  # type: ignore[arg-type]
            fmt="png",
            thread_count=2,
        )
        logger.info(f"[ocr] Rendered {len(images)} page image(s)")
    except Exception as e:
        logger.error(f"[ocr] pdf2image failed: {e}")
        raise RuntimeError(f"Could not render PDF pages to images: {e}")

    if not images:
        logger.error("[ocr] No images rendered — aborting")
        raise RuntimeError("PDF rendered 0 images.")

    # Step 2: OCR each page
    total_chars = 0
    for page_num, raw_image in enumerate(images):
        logger.info(f"[ocr] Processing page {page_num+1}/{len(images)}")

        if page_num > 0:
            doc.add_page_break()

        heading = doc.add_paragraph(f"— Page {page_num + 1} —")
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(10)
        heading.paragraph_format.space_after = Pt(6)

        # Preprocess
        try:
            processed = preprocess_image(raw_image, page_num + 1)
        except Exception as e:
            logger.warning(f"[ocr] Page {page_num+1}: preprocessing failed ({e}), using raw image")
            processed = raw_image

        # OCR attempt 1: preprocessed image
        try:
            ocr_text = pytesseract.image_to_string(processed, lang="eng", config=TESS_CONFIG)
            char_count = len(ocr_text.strip())
            logger.debug(f"[ocr] Page {page_num+1}: preprocessed OCR → {char_count} chars")
        except Exception as e:
            logger.error(f"[ocr] Page {page_num+1}: OCR on preprocessed image failed: {e}")
            ocr_text = ""
            char_count = 0

        # OCR attempt 2: fallback to raw image if preprocessed gave nothing
        if char_count == 0:
            logger.warning(f"[ocr] Page {page_num+1}: empty result — retrying with raw image")
            try:
                ocr_text = pytesseract.image_to_string(raw_image, lang="eng", config=TESS_CONFIG)
                char_count = len(ocr_text.strip())
                logger.debug(f"[ocr] Page {page_num+1}: raw OCR → {char_count} chars")
            except Exception as e:
                logger.error(f"[ocr] Page {page_num+1}: raw OCR also failed: {e}")
                ocr_text = ""

        if not ocr_text.strip():
            logger.warning(f"[ocr] Page {page_num+1}: both OCR attempts returned empty")
            doc.add_paragraph(f"[Page {page_num + 1} — no text could be extracted]")
            continue

        # Write lines to docx
        lines_written = 0
        for line in ocr_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            p = doc.add_paragraph(stripped)
            p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(3)
            lines_written += 1

        total_chars += char_count
        logger.info(f"[ocr] Page {page_num+1}: wrote {lines_written} lines ({char_count} chars)")

    logger.info(f"[ocr] All pages done. Total chars extracted: {total_chars}")

    doc.save(output_path)

    # Verify output
    size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
    logger.info(f"[ocr] Saved: {output_path} ({size} bytes)")
    if size < 5000:
        logger.warning(f"[ocr] Output is very small ({size} bytes) — OCR may have failed silently")


# ── WORD → PDF (LibreOffice) ────────────────────────────
def convert_word_to_pdf(input_path: str, output_dir: str, soffice: str) -> str:
    """Run LibreOffice headless to convert .docx → .pdf."""
    logger.info(f"[word-to-pdf] Input:      {input_path}")
    logger.info(f"[word-to-pdf] Output dir: {output_dir}")
    logger.info(f"[word-to-pdf] soffice:    {soffice}")

    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_path],
        capture_output=True,
        text=True,
        timeout=60,
    )

    logger.debug(f"[word-to-pdf] returncode: {result.returncode}")
    logger.debug(f"[word-to-pdf] stdout: {result.stdout.strip()}")
    if result.stderr.strip():
        logger.warning(f"[word-to-pdf] stderr: {result.stderr.strip()}")

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr}")

    base_name   = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}.pdf")
    logger.info(f"[word-to-pdf] Expected output: {output_path}")

    if not os.path.exists(output_path):
        # List what IS in the output dir to help debug
        contents = os.listdir(output_dir)
        logger.error(f"[word-to-pdf] Expected file not found. Dir contains: {contents}")
        raise RuntimeError("LibreOffice conversion produced no output file.")

    size = os.path.getsize(output_path)
    logger.info(f"[word-to-pdf] Output written: {output_path} ({size} bytes)")
    return output_path